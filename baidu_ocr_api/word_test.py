from docx import Document
from docx.shared import Pt

# 简单的打开word，输入数据，关闭word
document = Document()
# 向word里增加段落
document.add_paragraph('hello')

document.save('test.docx')
# 在一个段落中增加文字
# document = Document()
paragraph = document.add_paragraph('Normal text, ')
# 增加文字
paragraph.add_run('add text')

# 设置word字体大小
style = document.styles['Normal']
font = style.font
font.size = Pt(10)

document.save('test.docx')